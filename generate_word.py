import re
import requests
import zlib
import base64
from docx import Document
from docx.shared import Inches
import os

import time

def render_mermaid(mermaid_code):
    """Renders Mermaid code to PNG using mermaid.ink."""
    try:
        # Mermaid.ink expects base64 encoding (not zlib)
        payload = base64.b64encode(mermaid_code.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{payload}?theme=dark"
        
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = requests.get(url, timeout=30)
                if response.status_code == 200:
                    return response.content
                else:
                    print(f"Error rendering mermaid ({response.status_code}): {response.text}")
                    return None
            except (requests.exceptions.ConnectionError, requests.exceptions.Timeout) as e:
                print(f"Connection error on attempt {attempt + 1}: {e}")
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)
                else:
                    return None
    except Exception as e:
        print(f"Exception during rendering: {e}")
        return None

def generate_word_from_md(md_file, output_docx):
    if not os.path.exists(md_file):
        print(f"Markdown file not found: {md_file}")
        return

    with open(md_file, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    doc.add_heading('FindX System UML Diagrams', 0)

    # Split content by markdown headers to identify sections
    sections = re.split(r'\n(## .*)\n', content)
    
    # Process intro text if it exists
    intro = sections[0].strip()
    if intro and not intro.startswith('##'):
        doc.add_paragraph(intro)

    for i in range(1, len(sections), 2):
        header_raw = sections[i].strip()
        header = header_raw.replace('## ', '')
        body = sections[i+1] if i+1 < len(sections) else ""
        
        doc.add_heading(header, level=1)
        
        # Extract Mermaid code
        mermaid_match = re.search(r'```mermaid\n(.*?)\n```', body, re.DOTALL)
        
        # Extract description (text before and after the code block in the section)
        description_parts = re.split(r'```mermaid.*?```', body, flags=re.DOTALL)
        for part in description_parts:
            text = part.strip()
            if text:
                doc.add_paragraph(text)
        
        if mermaid_match:
            mermaid_code = mermaid_match.group(1).strip()
            doc.add_paragraph("UML Mermaid Code:", style='Heading 2')
            p = doc.add_paragraph()
            run = p.add_run(mermaid_code)
            run.font.name = 'Courier New'

    doc.save(output_docx)
    print(f"Successfully generated: {output_docx}")

if __name__ == "__main__":
    md_path = r"C:\Users\91960\.gemini\antigravity\brain\8856719f-fc41-4dc9-83bb-0858a4857ff3\uml_diagrams.md.resolved"
    output_path = r"d:\clg\FindX\findx_uml_diagrams.docx"
    generate_word_from_md(md_path, output_path)
